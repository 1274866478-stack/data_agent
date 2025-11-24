"""
RAG Document Processing Service for Data Agent V4 Backend
Handles document text extraction, chunking, and metadata extraction
"""

import asyncio
import logging
import re
import hashlib
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
import tempfile
import os
from io import BytesIO

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing libraries not available")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing library not available")

# Text processing
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter

from ..models.rag_models import (
    DocumentChunk,
    DocumentStatus,
    ChunkingStrategy,
    VectorCollection
)
from ..core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """RAG document processing service"""

    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
        self.max_file_size_mb = settings.max_document_size_mb if hasattr(settings, 'max_document_size_mb') else 50
        self.chunk_size = 512  # Default chunk size
        self.overlap_ratio = 0.2  # Default overlap ratio

        # Initialize NLTK data
        self._initialize_nltk()

        logger.info("RAG文档处理服务初始化完成")

    def _initialize_nltk(self):
        """Initialize NLTK data for text processing"""
        try:
            # Download required NLTK data if not present
            nltk_data_path = os.path.join(tempfile.gettempdir(), 'nltk_data')
            os.makedirs(nltk_data_path, exist_ok=True)

            try:
                nltk.data.find('tokenizers/punkt', paths=[nltk_data_path])
            except LookupError:
                nltk.download('punkt', download_dir=nltk_data_path)

            try:
                nltk.data.find('corpora/stopwords', paths=[nltk_data_path])
            except LookupError:
                nltk.download('stopwords', download_dir=nltk_data_path)

        except Exception as e:
            logger.warning(f"NLTK初始化失败: {e}")

    async def process_document(
        self,
        file_content: bytes,
        file_name: str,
        file_type: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.SEMANTIC,
        chunk_size: int = 512,
        overlap_ratio: float = 0.2
    ) -> Tuple[List[DocumentChunk], Dict[str, Any]]:
        """
        Process uploaded document into chunks

        Args:
            file_content: File content as bytes
            file_name: Original file name
            file_type: File type (pdf, docx, txt)
            document_id: Document ID in database
            tenant_id: Tenant ID for isolation
            collection_id: Vector collection ID
            chunking_strategy: Chunking strategy to use
            chunk_size: Target chunk size in tokens
            overlap_ratio: Overlap ratio between chunks

        Returns:
            Tuple of (chunks list, processing metadata)
        """
        start_time = datetime.utcnow()

        try:
            # Validate inputs
            await self._validate_document(file_content, file_name, file_type)

            # Extract text content
            logger.info(f"开始提取文档文本: {file_name}")
            text_content, metadata = await self._extract_text(file_content, file_type, file_name)

            if not text_content.strip():
                raise ValueError("文档文本提取失败或文档为空")

            # Clean and preprocess text
            logger.info(f"清理和预处理文本，长度: {len(text_content)}")
            cleaned_text = await self._clean_text(text_content)

            # Create chunks based on strategy
            logger.info(f"使用{chunking_strategy.value}策略创建文档块")
            chunks = await self._create_chunks(
                cleaned_text,
                document_id,
                tenant_id,
                collection_id,
                file_name,
                file_type,
                chunking_strategy,
                chunk_size,
                overlap_ratio,
                metadata
            )

            # Calculate processing metadata
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            processing_metadata = {
                'file_name': file_name,
                'file_type': file_type,
                'file_size_mb': len(file_content) / (1024 * 1024),
                'extracted_characters': len(text_content),
                'cleaned_characters': len(cleaned_text),
                'chunks_created': len(chunks),
                'chunking_strategy': chunking_strategy.value,
                'chunk_size': chunk_size,
                'overlap_ratio': overlap_ratio,
                'processing_time_ms': processing_time,
                'extraction_metadata': metadata,
                'quality_metrics': await self._calculate_document_quality(cleaned_text, chunks)
            }

            logger.info(f"文档处理完成: {file_name} -> {len(chunks)}个块，耗时{processing_time:.1f}ms")
            return chunks, processing_metadata

        except Exception as e:
            error_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            logger.error(f"文档处理失败 {file_name}: {str(e)}")

            # Return error metadata
            error_metadata = {
                'file_name': file_name,
                'file_type': file_type,
                'error_message': str(e),
                'processing_time_ms': error_time,
                'status': 'failed'
            }

            raise RuntimeError(f"文档处理失败: {str(e)}") from e

    async def _validate_document(
        self,
        file_content: bytes,
        file_name: str,
        file_type: str
    ):
        """Validate document before processing"""
        # Check file size
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            raise ValueError(f"文件大小({file_size_mb:.1f}MB)超过限制({self.max_file_size_mb}MB)")

        # Check file type
        if file_type.lower() not in self.supported_formats:
            raise ValueError(f"不支持的文件类型: {file_type}. 支持的格式: {', '.join(self.supported_formats)}")

        # Check if file is empty
        if len(file_content) == 0:
            raise ValueError("文件为空")

        # Validate file content by magic bytes
        await self._validate_file_content(file_content, file_type)

    async def _validate_file_content(self, file_content: bytes, file_type: str):
        """Validate file content by checking magic bytes"""
        file_type = file_type.lower()

        if file_type == 'pdf':
            # PDF magic bytes: %PDF-
            if not file_content.startswith(b'%PDF-'):
                raise ValueError("无效的PDF文件格式")
        elif file_type == 'docx':
            # DOCX magic bytes: PK (zip archive)
            if not file_content.startswith(b'PK'):
                raise ValueError("无效的DOCX文件格式")
        elif file_type == 'txt':
            # Text file - check if it's actually text
            try:
                file_content.decode('utf-8')
            except UnicodeDecodeError:
                raise ValueError("文本文件编码无效，请使用UTF-8编码")

    async def _extract_text(
        self,
        file_content: bytes,
        file_type: str,
        file_name: str
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from document"""
        file_type = file_type.lower()

        if file_type == 'pdf':
            return await self._extract_pdf_text(file_content, file_name)
        elif file_type == 'docx':
            return await self._extract_docx_text(file_content, file_name)
        elif file_type == 'txt':
            return await self._extract_txt_text(file_content, file_name)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    async def _extract_pdf_text(self, file_content: bytes, file_name: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pdfplumber (primary) and PyPDF2 (fallback)"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF处理库未安装，请安装 PyPDF2 和 pdfplumber")

        metadata = {
            'extraction_method': 'pdfplumber',
            'pages_processed': 0,
            'has_images': False,
            'has_tables': False
        }

        text_content = ""

        try:
            # Primary method: pdfplumber
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                metadata['pages_processed'] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        # Add page separator with metadata
                        text_content += f"\n\n--- 第{page_num}页 ---\n\n"
                        text_content += page_text

                        # Check for tables
                        tables = page.extract_tables()
                        if tables:
                            metadata['has_tables'] = True
                            for table in tables:
                                table_text = self._format_table_text(table)
                                if table_text:
                                    text_content += f"\n\n表格:\n{table_text}\n"

                        # Check for images (pdfplumber can detect images)
                        if hasattr(page, 'images') and page.images:
                            metadata['has_images'] = True

        except Exception as e:
            logger.warning(f"pdfplumber提取失败，尝试PyPDF2: {e}")

            # Fallback method: PyPDF2
            try:
                import PyPDF2
                from PyPDF2 import PdfReader

                reader = PdfReader(BytesIO(file_content))
                metadata['extraction_method'] = 'PyPDF2'
                metadata['pages_processed'] = len(reader.pages)

                text_content = ""
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n\n--- 第{page_num}页 ---\n\n"
                        text_content += page_text

            except Exception as fallback_error:
                raise RuntimeError(f"PDF文本提取失败: {str(fallback_error)}")

        return text_content.strip(), metadata

    async def _extract_docx_text(self, file_content: bytes, file_name: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX document"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX处理库未安装，请安装 python-docx")

        metadata = {
            'extraction_method': 'python-docx',
            'paragraphs_processed': 0,
            'has_tables': False,
            'has_images': False,
            'document_structure': []
        }

        try:
            doc = Document(BytesIO(file_content))
            text_content = ""

            # Process paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"
                    metadata['paragraphs_processed'] += 1

                    # Track document structure
                    if paragraph.style.name.startswith('Heading'):
                        metadata['document_structure'].append({
                            'type': 'heading',
                            'style': paragraph.style.name,
                            'text': paragraph.text.strip()[:100],
                            'paragraph_index': i
                        })

            # Process tables
            for table in doc.tables:
                metadata['has_tables'] = True
                table_text = self._format_docx_table(table)
                if table_text:
                    text_content += f"\n\n表格:\n{table_text}\n"

            # Check for images (basic detection)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata['has_images'] = True
                    break

            return text_content.strip(), metadata

        except Exception as e:
            raise RuntimeError(f"DOCX文本提取失败: {str(e)}")

    async def _extract_txt_text(self, file_content: bytes, file_name: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        metadata = {
            'extraction_method': 'utf-8-decode',
            'encoding': 'utf-8',
            'line_count': 0,
            'character_count': 0
        }

        try:
            # Try UTF-8 first
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Try with error handling
                text_content = file_content.decode('utf-8', errors='replace')
                metadata['encoding'] = 'utf-8-with-replacement'
            except Exception as e:
                # Last resort: try latin-1
                text_content = file_content.decode('latin-1', errors='replace')
                metadata['encoding'] = 'latin-1-with-replacement'

        # Update metadata
        lines = text_content.split('\n')
        metadata['line_count'] = len(lines)
        metadata['character_count'] = len(text_content)

        return text_content, metadata

    def _format_table_text(self, table: List[List[str]]) -> str:
        """Format table data into readable text"""
        if not table:
            return ""

        formatted_lines = []

        for row_num, row in enumerate(table):
            if not row:
                continue

            # Clean cell content and format as table row
            clean_cells = [str(cell).strip() if cell else "" for cell in row]
            formatted_row = " | ".join(clean_cells)
            formatted_lines.append(formatted_row)

            # Add separator after header row
            if row_num == 0:
                separator = "-" * len(formatted_row)
                formatted_lines.append(separator)

        return "\n".join(formatted_lines)

    def _format_docx_table(self, table) -> str:
        """Format python-docx table into readable text"""
        if not table.rows:
            return ""

        formatted_lines = []

        for row_num, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                row_cells.append(cell.text.strip())

            if any(row_cells):  # Skip empty rows
                formatted_row = " | ".join(row_cells)
                formatted_lines.append(formatted_row)

                # Add separator after header row
                if row_num == 0:
                    separator = "-" * len(formatted_row)
                    formatted_lines.append(separator)

        return "\n".join(formatted_lines)

    async def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,!?;:()[\]{}"\'-]', ' ', text)

        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)

        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)

        # Clean up spaces
        text = re.sub(r' [ ]+', ' ', text)

        return text.strip()

    async def _create_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunking_strategy: ChunkingStrategy,
        chunk_size: int,
        overlap_ratio: float,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create document chunks based on strategy"""
        if chunking_strategy == ChunkingStrategy.SEMANTIC:
            return await self._create_semantic_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, overlap_ratio, metadata
            )
        elif chunking_strategy == ChunkingStrategy.FIXED_SIZE:
            return await self._create_fixed_size_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, overlap_ratio, metadata
            )
        elif chunking_strategy == ChunkingStrategy.OVERLAP:
            return await self._create_overlap_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, overlap_ratio, metadata
            )
        elif chunking_strategy == ChunkingStrategy.PARAGRAPH:
            return await self._create_paragraph_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, metadata
            )
        elif chunking_strategy == ChunkingStrategy.SENTENCE:
            return await self._create_sentence_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, overlap_ratio, metadata
            )
        else:
            # Default to semantic chunking
            return await self._create_semantic_chunks(
                text, document_id, tenant_id, collection_id, file_name, file_type, chunk_size, overlap_ratio, metadata
            )

    async def _create_semantic_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunk_size: int,
        overlap_ratio: float,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create chunks using semantic analysis"""
        chunks = []

        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        current_chunk = ""
        current_chunk_index = 0
        overlap_size = int(chunk_size * overlap_ratio)

        for paragraph in paragraphs:
            # Check paragraph size
            paragraph_tokens = self._estimate_tokens(paragraph)

            if paragraph_tokens > chunk_size * 1.5:
                # Large paragraph, split it further
                sub_chunks = await self._split_large_paragraph(
                    paragraph, chunk_size, overlap_size
                )

                for sub_chunk in sub_chunks:
                    if sub_chunk.strip():
                        chunk = await self._create_document_chunk(
                            sub_chunk.strip(),
                            current_chunk_index,
                            document_id,
                            tenant_id,
                            collection_id,
                            file_name,
                            file_type,
                            ChunkingStrategy.SEMANTIC,
                            chunk_size,
                            0,  # Overlap handled at sub-chunk level
                            metadata
                        )
                        chunks.append(chunk)
                        current_chunk_index += 1

            else:
                # Check if adding this paragraph would exceed chunk size
                current_chunk_tokens = self._estimate_tokens(current_chunk)

                if current_chunk_tokens + paragraph_tokens <= chunk_size:
                    # Add to current chunk
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph
                else:
                    # Save current chunk and start new one
                    if current_chunk.strip():
                        chunk = await self._create_document_chunk(
                            current_chunk.strip(),
                            current_chunk_index,
                            document_id,
                            tenant_id,
                            collection_id,
                            file_name,
                            file_type,
                            ChunkingStrategy.SEMANTIC,
                            chunk_size,
                            0,
                            metadata
                        )
                        chunks.append(chunk)
                        current_chunk_index += 1

                    current_chunk = paragraph

        # Don't forget the last chunk
        if current_chunk.strip():
            chunk = await self._create_document_chunk(
                current_chunk.strip(),
                current_chunk_index,
                document_id,
                tenant_id,
                collection_id,
                file_name,
                file_type,
                ChunkingStrategy.SEMANTIC,
                chunk_size,
                0,
                metadata
            )
            chunks.append(chunk)

        return chunks

    async def _create_fixed_size_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunk_size: int,
        overlap_ratio: float,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create chunks with fixed token size"""
        chunks = []
        words = text.split()
        overlap_size = int(chunk_size * overlap_ratio)

        current_chunk_index = 0
        i = 0

        while i < len(words):
            # Determine chunk boundaries
            start = i
            end = min(i + chunk_size, len(words))

            # Extract chunk words
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            if chunk_text.strip():
                chunk = await self._create_document_chunk(
                    chunk_text.strip(),
                    current_chunk_index,
                    document_id,
                    tenant_id,
                    collection_id,
                    file_name,
                    file_type,
                    ChunkingStrategy.FIXED_SIZE,
                    chunk_size,
                    overlap_size if end < len(words) else 0,
                    metadata
                )
                chunks.append(chunk)
                current_chunk_index += 1

            # Move to next chunk with overlap
            i = end - overlap_size if end < len(words) else end

        return chunks

    async def _create_overlap_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunk_size: int,
        overlap_ratio: float,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create chunks with overlapping content"""
        chunks = []
        words = text.split()
        overlap_size = int(chunk_size * overlap_ratio)

        current_chunk_index = 0
        i = 0

        while i < len(words):
            # Determine chunk boundaries
            start = max(0, i - overlap_size) if i > 0 else 0
            end = min(start + chunk_size, len(words))

            # Adjust start if we're at the end
            if end == len(words) and end - start < chunk_size:
                start = max(0, end - chunk_size)

            # Extract chunk words
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            if chunk_text.strip():
                chunk = await self._create_document_chunk(
                    chunk_text.strip(),
                    current_chunk_index,
                    document_id,
                    tenant_id,
                    collection_id,
                    file_name,
                    file_type,
                    ChunkingStrategy.OVERLAP,
                    chunk_size,
                    overlap_size,
                    metadata
                )
                chunks.append(chunk)
                current_chunk_index += 1

            # Move to next chunk
            i = end - overlap_size if end < len(words) else end

        return chunks

    async def _create_paragraph_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunk_size: int,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create chunks based on paragraph boundaries"""
        chunks = []
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        current_chunk_index = 0

        for paragraph in paragraphs:
            # Check if paragraph is too large
            if self._estimate_tokens(paragraph) > chunk_size:
                # Split large paragraph
                sub_chunks = await self._split_large_paragraph(paragraph, chunk_size, 0)

                for sub_chunk in sub_chunks:
                    if sub_chunk.strip():
                        chunk = await self._create_document_chunk(
                            sub_chunk.strip(),
                            current_chunk_index,
                            document_id,
                            tenant_id,
                            collection_id,
                            file_name,
                            file_type,
                            ChunkingStrategy.PARAGRAPH,
                            chunk_size,
                            0,
                            metadata
                        )
                        chunks.append(chunk)
                        current_chunk_index += 1
            else:
                # Use paragraph as is
                chunk = await self._create_document_chunk(
                    paragraph,
                    current_chunk_index,
                    document_id,
                    tenant_id,
                    collection_id,
                    file_name,
                    file_type,
                    ChunkingStrategy.PARAGRAPH,
                    chunk_size,
                    0,
                    metadata
                )
                chunks.append(chunk)
                current_chunk_index += 1

        return chunks

    async def _create_sentence_chunks(
        self,
        text: str,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunk_size: int,
        overlap_ratio: float,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Create chunks based on sentence boundaries"""
        chunks = []

        try:
            # Use NLTK sentence tokenizer
            sentences = sent_tokenize(text)
        except Exception:
            # Fallback: simple sentence splitting
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if s.strip()]

        current_chunk = ""
        current_chunk_index = 0
        overlap_size = int(chunk_size * overlap_ratio)

        for sentence in sentences:
            sentence_tokens = self._estimate_tokens(sentence)
            current_chunk_tokens = self._estimate_tokens(current_chunk)

            if current_chunk_tokens + sentence_tokens <= chunk_size:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
            else:
                # Save current chunk and start new one
                if current_chunk.strip():
                    chunk = await self._create_document_chunk(
                        current_chunk.strip(),
                        current_chunk_index,
                        document_id,
                        tenant_id,
                        collection_id,
                        file_name,
                        file_type,
                        ChunkingStrategy.SENTENCE,
                        chunk_size,
                        0,
                        metadata
                    )
                    chunks.append(chunk)
                    current_chunk_index += 1

                # Handle very long sentences
                if sentence_tokens > chunk_size:
                    sub_chunks = await self._split_large_sentence(sentence, chunk_size)
                    for sub_chunk in sub_chunks:
                        if sub_chunk.strip():
                            chunk = await self._create_document_chunk(
                                sub_chunk.strip(),
                                current_chunk_index,
                                document_id,
                                tenant_id,
                                collection_id,
                                file_name,
                                file_type,
                                ChunkingStrategy.SENTENCE,
                                chunk_size,
                                0,
                                metadata
                            )
                            chunks.append(chunk)
                            current_chunk_index += 1

                    current_chunk = ""
                else:
                    current_chunk = sentence

        # Don't forget the last chunk
        if current_chunk.strip():
            chunk = await self._create_document_chunk(
                current_chunk.strip(),
                current_chunk_index,
                document_id,
                tenant_id,
                collection_id,
                file_name,
                file_type,
                ChunkingStrategy.SENTENCE,
                chunk_size,
                0,
                metadata
            )
            chunks.append(chunk)

        return chunks

    async def _split_large_paragraph(
        self,
        paragraph: str,
        chunk_size: int,
        overlap_size: int
    ) -> List[str]:
        """Split a large paragraph into smaller chunks"""
        chunks = []
        sentences = re.split(r'[.!?]+', paragraph)
        sentences = [s.strip() for s in sentences if s.strip()]

        current_chunk = ""
        current_tokens = 0

        for sentence in sentences:
            sentence_tokens = self._estimate_tokens(sentence)

            if current_tokens + sentence_tokens <= chunk_size:
                current_chunk += sentence + ". "
                current_tokens += sentence_tokens
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # Handle very long sentences
                if sentence_tokens > chunk_size:
                    words = sentence.split()
                    word_chunks = []

                    for i in range(0, len(words), chunk_size):
                        word_chunk = " ".join(words[i:i + chunk_size])
                        word_chunks.append(word_chunk)

                    chunks.extend(word_chunks)
                    current_chunk = ""
                    current_tokens = 0
                else:
                    current_chunk = sentence + ". "
                    current_tokens = sentence_tokens

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    async def _split_large_sentence(self, sentence: str, chunk_size: int) -> List[str]:
        """Split a large sentence into smaller chunks"""
        words = sentence.split()
        chunks = []

        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk = " ".join(chunk_words)
            if chunk.strip():
                chunks.append(chunk.strip())

        return chunks

    async def _create_document_chunk(
        self,
        content: str,
        chunk_index: int,
        document_id: int,
        tenant_id: str,
        collection_id: str,
        file_name: str,
        file_type: str,
        chunking_strategy: ChunkingStrategy,
        chunk_size: int,
        overlap_size: int,
        metadata: Dict[str, Any]
    ) -> DocumentChunk:
        """Create a DocumentChunk object"""

        # Extract page number and chapter from content if available
        page_number = self._extract_page_number(content)
        chapter_title = self._extract_chapter_title(content)

        # Calculate quality score
        quality_score = await self._calculate_chunk_quality(content)

        return DocumentChunk(
            document_id=document_id,
            tenant_id=tenant_id,
            chunk_index=chunk_index,
            content=content,
            content_length=len(content),
            token_count=self._estimate_tokens(content),
            file_name=file_name,
            file_type=file_type,
            page_number=page_number,
            chapter_title=chapter_title,
            chunking_strategy=chunking_strategy,
            chunk_size=chunk_size,
            overlap_size=overlap_size,
            quality_score=quality_score,
            processed_at=datetime.utcnow()
        )

    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count for text (rough approximation)"""
        if not text:
            return 0

        # Simple heuristic: ~4 characters per token for English, ~1.5 for Chinese
        english_chars = len(re.findall(r'[a-zA-Z\s]', text))
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))

        english_tokens = english_chars / 4
        chinese_tokens = chinese_chars / 1.5

        return int(english_tokens + chinese_tokens)

    def _extract_page_number(self, content: str) -> Optional[int]:
        """Extract page number from content"""
        # Look for page number patterns
        page_patterns = [
            r'--- 第(\d+)页 ---',
            r'第\s*(\d+)\s*页',
            r'Page\s*(\d+)',
            r'(\d+)\s*/\s*\d+',  # e.g., "5/23"
        ]

        for pattern in page_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                try:
                    return int(match.group(1))
                except ValueError:
                    continue

        return None

    def _extract_chapter_title(self, content: str) -> Optional[str]:
        """Extract chapter or section title from content"""
        # Look for heading patterns
        heading_patterns = [
            r'^(第[一二三四五六七八九十\d]+章[：:]\s*.+)',
            r'^(Chapter\s*\d+[：:]\s*.+)',
            r'^(第\d+节[：:]\s*.+)',
            r'^([一二三四五六七八九十\d]+[、\.]\s*.+)',
            r'^(\d+\.\s*.+)',
        ]

        lines = content.split('\n')
        for line in lines[:3]:  # Check first few lines
            line = line.strip()
            if len(line) < 100 and len(line) > 5:  # Reasonable title length
                for pattern in heading_patterns:
                    if re.match(pattern, line):
                        return line

        return None

    async def _calculate_chunk_quality(self, content: str) -> float:
        """Calculate quality score for a chunk (0-1)"""
        if not content:
            return 0.0

        score = 0.0

        # Length score (optimal between 100-1000 characters)
        length = len(content)
        if 100 <= length <= 1000:
            score += 0.3
        elif 50 <= length <= 2000:
            score += 0.2
        else:
            score += 0.1

        # Content diversity (presence of different word types)
        words = content.lower().split()
        if words:
            unique_words = len(set(words))
            diversity = unique_words / len(words)
            score += diversity * 0.2

        # Punctuation score (well-structured text)
        punctuation_chars = len(re.findall(r'[.,!?;:]', content))
        if words:
            punctuation_ratio = punctuation_chars / len(words)
            if 0.05 <= punctuation_ratio <= 0.2:
                score += 0.2
            elif punctuation_ratio > 0:
                score += 0.1

        # Alphanumeric content
        alphanumeric_ratio = len(re.findall(r'[a-zA-Z\u4e00-\u9fff0-9]', content)) / len(content) if content else 0
        score += alphanumeric_ratio * 0.3

        return min(score, 1.0)

    async def _calculate_document_quality(
        self,
        text: str,
        chunks: List[DocumentChunk]
    ) -> Dict[str, Any]:
        """Calculate overall document quality metrics"""
        return {
            'total_characters': len(text),
            'total_chunks': len(chunks),
            'average_chunk_length': sum(c.content_length for c in chunks) / len(chunks) if chunks else 0,
            'average_quality_score': sum(c.quality_score or 0 for c in chunks) / len(chunks) if chunks else 0,
            'content_languages': self._detect_languages(text),
            'readability_score': self._calculate_readability_score(text),
            'content_type': self._detect_content_type(text)
        }

    def _detect_languages(self, text: str) -> List[str]:
        """Detect languages present in text"""
        languages = []

        # Simple language detection based on character ranges
        if re.search(r'[\u4e00-\u9fff]', text):
            languages.append('chinese')
        if re.search(r'[a-zA-Z]', text):
            languages.append('english')

        return languages if languages else ['unknown']

    def _calculate_readability_score(self, text: str) -> float:
        """Calculate simple readability score (0-1)"""
        if not text:
            return 0.0

        words = text.split()
        sentences = re.split(r'[.!?]+', text)

        if not words or not sentences:
            return 0.0

        avg_words_per_sentence = len(words) / len(sentences)

        # Simple readability: prefer 15-25 words per sentence
        if 15 <= avg_words_per_sentence <= 25:
            return 1.0
        elif 10 <= avg_words_per_sentence <= 35:
            return 0.8
        elif 5 <= avg_words_per_sentence <= 50:
            return 0.6
        else:
            return 0.4

    def _detect_content_type(self, text: str) -> str:
        """Detect the type of content (simple heuristic)"""
        text_lower = text.lower()

        if any(keyword in text_lower for keyword in ['摘要', 'abstract', '总结', 'conclusion']):
            return 'academic'
        elif any(keyword in text_lower for keyword in ['条款', 'contract', '协议', 'agreement']):
            return 'legal'
        elif any(keyword in text_lower for keyword in ['收入', 'revenue', '利润', 'profit', '财务', 'financial']):
            return 'business'
        elif any(keyword in text_lower for keyword in ['方法', 'method', '步骤', 'procedure', '指南', 'guide']):
            return 'technical'
        else:
            return 'general'

    async def get_processing_status(
        self,
        document_id: int,
        tenant_id: str
    ) -> Dict[str, Any]:
        """Get processing status for a document"""
        # This would typically query the database for processing status
        # For now, return a placeholder implementation
        return {
            'document_id': document_id,
            'tenant_id': tenant_id,
            'status': 'unknown',
            'progress_percentage': 0,
            'current_step': 'unknown',
            'estimated_time_remaining': None,
            'error_message': None
        }

    async def cleanup_temporary_resources(self):
        """Clean up temporary resources used during processing"""
        # Clean up NLTK data if needed
        pass